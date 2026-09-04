# VetarAI - Local-first multi-agent orchestration application
# Copyright (C) 2026 zero11924065-dev
#
# This file is part of VetarAI.
#
# VetarAI is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# VetarAI is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE. See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with VetarAI. If not, see <https://www.gnu.org/licenses/>.
"""M7（TS-113）：圆桌附件内置解析器（核心，非插件）。

支持格式（3.17 第三项）：
- PDF（pypdf 逐页）/ Word .docx（python-docx 段落+表格）/
  Excel .xlsx（openpyxl 逐 sheet 逐行）/ CSV（标准库）/ 纯文本族（utf-8/gbk）
- 图片（.png/.jpg/.jpeg/.gif/.webp）→ 可选视觉模型识别（vision_parse_attachments 开关）
- 其余格式 → 返回 None（调用方仅标注，与现状一致）

约定：解析失败一律返回 None（调用方标注，不阻塞圆桌创建）；
截断由调用方按单文件/总量限制处理。
"""
from __future__ import annotations

import csv
import io

# 文本族扩展名（小写，含点）
TEXT_EXTS = {".txt", ".md", ".markdown", ".json", ".yaml", ".yml", ".ini",
             ".log", ".py", ".js", ".ts", ".tsx", ".jsx", ".html", ".htm",
             ".xml", ".toml", ".cfg", ".conf", ".sh", ".css", ".csv"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
SUPPORTED_EXTS = TEXT_EXTS | IMAGE_EXTS | {".pdf", ".docx", ".xlsx", ".xlsm"}

# Excel 防爆炸：单 sheet 最多行数 / 单元格截断
_XLSX_MAX_ROWS_PER_SHEET = 200
_XLSX_MAX_CELL_LEN = 200


def _ext_of(name: str) -> str:
    n = str(name or "").lower()
    return n[n.rfind("."):] if "." in n else ""


def _parse_text(raw: bytes) -> str | None:
    for enc in ("utf-8", "gbk", "latin-1"):
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return None


def _parse_csv(raw: bytes) -> str | None:
    text = _parse_text(raw)
    if text is None:
        return None
    try:
        rows = list(csv.reader(io.StringIO(text)))[:500]
        lines = [" | ".join(str(c)[:_XLSX_MAX_CELL_LEN] for c in row) for row in rows if row]
        return "\n".join(lines) if lines else None
    except Exception:
        return None


def _parse_pdf(raw: bytes) -> str | None:
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(raw))
        parts = []
        for i, page in enumerate(reader.pages):
            t = (page.extract_text() or "").strip()
            if t:
                parts.append(f"[第{i + 1}页]\n{t}")
        return "\n\n".join(parts) if parts else None
    except Exception:
        return None


def _parse_docx(raw: bytes) -> str | None:
    try:
        import docx
        d = docx.Document(io.BytesIO(raw))
        parts: list[str] = []
        for para in d.paragraphs:
            t = (para.text or "").strip()
            if t:
                parts.append(t)
        for table in d.tables:
            for row in table.rows:
                cells = [(c.text or "").strip().replace("\n", " ") for c in row.cells]
                line = " | ".join(c for c in cells if c)
                if line.strip():
                    parts.append(line)
        return "\n".join(parts) if parts else None
    except Exception:
        return None


def _parse_xlsx(raw: bytes) -> str | None:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
        parts: list[str] = []
        for ws in wb.worksheets:
            lines = [f"[工作表: {ws.title}]"]
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= _XLSX_MAX_ROWS_PER_SHEET:
                    lines.append(f"（后续行已省略，共超 {i} 行）")
                    break
                vals = [str(c)[:_XLSX_MAX_CELL_LEN] if c is not None else "" for c in row]
                line = " | ".join(v for v in vals if v != "")
                if line.strip():
                    lines.append(line)
            if len(lines) > 1:
                parts.append("\n".join(lines))
        wb.close()
        return "\n\n".join(parts) if parts else None
    except Exception:
        return None


def parse_attachment(name: str, raw: bytes) -> tuple[str | None, str]:
    """解析附件内容为文本。

    返回 (text, kind)：
    - text: 解析出的文本；无法解析/失败 → None
    - kind: "text"/"pdf"/"docx"/"xlsx"/"csv"/"image"/"binary"（标注用）

    图片：本函数不识别（返回 (None, "image")）；视觉识别由异步调用方
    （app.py，因连接器 chat 为异步）在配置开关开启时单独完成。
    """
    ext = _ext_of(name)
    if ext == ".pdf":
        return _parse_pdf(raw), "pdf"
    if ext == ".docx":
        return _parse_docx(raw), "docx"
    if ext in (".xlsx", ".xlsm"):
        return _parse_xlsx(raw), "xlsx"
    if ext == ".csv":
        return _parse_csv(raw), "csv"
    if ext in TEXT_EXTS:
        return _parse_text(raw), "text"
    if ext in IMAGE_EXTS:
        return None, "image"
    return None, "binary"
